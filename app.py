import streamlit as st
import os
import glob
from docx import Document
import re
from datetime import datetime
import requests
from io import BytesIO

def extract_company_name_from_filename(filename):
    """Extract company name from filename like 'bajaj-auto_complete_20250819_172757.docx'"""
    # Remove file extension
    name = os.path.splitext(filename)[0]
    
    # Remove timestamp pattern (e.g., _complete_20250819_172757)
    name = re.sub(r'_complete_\d{8}_\d{6}$', '', name)
    
    # Replace hyphens and underscores with spaces and title case
    company_name = name.replace('-', ' ').replace('_', ' ').title()
    
    return company_name

def extract_data_from_docx_content(content):
    """Extract company data from DOCX text content"""
    try:
        # Extract summary info
        summary_match = re.search(r'\*\*Summary:\*\*\s*(.+?)(?:\n|\r|$)', content)
        summary = summary_match.group(1).strip() if summary_match else "No summary available"
        
        # Extract date from summary
        date_match = re.search(r'Date:\s*(\d{4}-\d{2}-\d{2}\s+\d{2}:\d{2}:\d{2})', summary)
        extraction_date = date_match.group(1) if date_match else "Unknown date"
        
        # Extract total links count from summary
        links_count_match = re.search(r'Total links:\s*(\d+)', summary)
        total_links_from_summary = int(links_count_match.group(1)) if links_count_match else 0
        
        # Extract text length from summary
        text_length_match = re.search(r'Text length:\s*(\d+(?:,\d+)*)', summary)
        text_length_from_summary = int(text_length_match.group(1).replace(',', '')) if text_length_match else 0
        
        # Extract links section
        links = []
        link_section_started = False
        text_section_started = False
        
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            
            # Check for section headers
            if '# Extracted Links' in line:
                link_section_started = True
                text_section_started = False
                continue
            elif '# Extracted Text Content' in line:
                link_section_started = False
                text_section_started = True
                continue
            
            # Extract links
            if link_section_started and line.startswith('https://'):
                links.append(line)
        
        # Extract text content section
        text_content_lines = []
        text_section_started = False
        
        for line in lines:
            if '# Extracted Text Content' in line:
                text_section_started = True
                continue
            
            if text_section_started and line.strip():
                # Clean up the line (remove backslashes and extra formatting)
                cleaned_line = line.replace('\\', '').strip()
                if cleaned_line:
                    text_content_lines.append(cleaned_line)
        
        text_content = '\n\n'.join(text_content_lines)
        
        return {
            'summary': summary,
            'extraction_date': extraction_date,
            'links': links,
            'text_content': text_content,
            'total_links': len(links),
            'total_links_from_summary': total_links_from_summary,
            'text_length': len(text_content),
            'text_length_from_summary': text_length_from_summary
        }
        
    except Exception as e:
        st.error(f"Error processing content: {str(e)}")
        return None

def extract_data_from_docx_file(file_path):
    """Extract company data from DOCX file"""
    try:
        # Read the document
        if isinstance(file_path, str):
            doc = Document(file_path)
        else:
            # If it's a BytesIO object (from URL)
            doc = Document(file_path)
        
        # Get full text
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)
        
        content = '\n'.join(full_text)
        return extract_data_from_docx_content(content)
        
    except Exception as e:
        st.error(f"Error processing file: {str(e)}")
        return None

def load_docx_files_from_directory(directory_path):
    """Load all DOCX files from a directory"""
    docx_files = glob.glob(os.path.join(directory_path, "*.docx"))
    company_data = []
    
    for file_path in docx_files:
        filename = os.path.basename(file_path)
        company_name = extract_company_name_from_filename(filename)
        
        data = extract_data_from_docx_file(file_path)
        if data:
            company_data.append({
                'company_name': company_name,
                'filename': filename,
                'file_path': file_path,
                **data
            })
    
    return company_data

def load_docx_from_github(github_raw_urls):
    """Load DOCX files from GitHub raw URLs"""
    company_data = []
    
    for url in github_raw_urls:
        try:
            # Extract filename from URL
            filename = url.split('/')[-1]
            company_name = extract_company_name_from_filename(filename)
            
            st.write(f"🔄 Processing {filename}...")
            
            # Download file with proper headers
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
            response = requests.get(url, headers=headers, timeout=30)
            response.raise_for_status()
            
            st.write(f"✅ Downloaded {filename} ({len(response.content)} bytes)")
            
            # Create BytesIO object
            docx_file = BytesIO(response.content)
            
            # Extract data
            data = extract_data_from_docx_file(docx_file)
            if data:
                company_data.append({
                    'company_name': company_name,
                    'filename': filename,
                    'github_url': url,
                    **data
                })
                st.success(f"✅ Successfully processed {company_name}")
            else:
                st.warning(f"⚠️ No data extracted from {filename}")
                
        except requests.exceptions.RequestException as e:
            st.error(f"❌ Network error loading {filename}: {str(e)}")
        except Exception as e:
            st.error(f"❌ Error processing {filename}: {str(e)}")
            st.write(f"URL: {url}")
    
    return company_data

def display_company_card(company_info):
    """Display a comprehensive card for each company with full content"""
    
    # Create an expandable card with company name and key metrics
    with st.expander(
        f"🏢 **{company_info['company_name']}** | 🔗 {company_info['total_links']} Links | 📝 {company_info['text_length']:,} chars", 
        expanded=True  # Keep expanded to show full content by default
    ):
        
        # Header info
        col1, col2 = st.columns([2, 1])
        with col1:
            st.markdown(f"**📄 File:** `{company_info['filename']}`")
            if 'extraction_date' in company_info:
                st.markdown(f"**📅 Extracted:** {company_info['extraction_date']}")
        with col2:
            # Validation check
            if company_info['total_links'] != company_info.get('total_links_from_summary', 0):
                st.warning(f"⚠️ Link count mismatch: Found {company_info['total_links']}, Expected {company_info.get('total_links_from_summary', 0)}")
        
        st.markdown("---")
        
        # Summary information
        st.markdown("### 📋 Document Summary")
        st.info(company_info['summary'])
        
        # Links section - SHOW ALL LINKS
        st.markdown(f"### 🔗 All Extracted Links ({company_info['total_links']} total)")
        
        if company_info['links']:
            # Display all links in a clean format
            for i, link in enumerate(company_info['links'], 1):
                # Extract domain for better display
                domain = re.search(r'https?://(?:www\.)?([^/]+)', link)
                domain_name = domain.group(1) if domain else "Unknown"
                
                # Show full link with index
                st.markdown(f"**{i}.** [{domain_name}]({link})")
                st.caption(f"Full URL: {link}")
                
                # Add some spacing every 5 links for readability
                if i % 5 == 0 and i < len(company_info['links']):
                    st.markdown("---")
        else:
            st.warning("No links found")
        
        st.markdown("---")
        
        # Text content section - SHOW COMPLETE TEXT
        st.markdown("### 📖 Complete Extracted Text Content")
        if company_info['text_content']:
            # Show character count
            st.caption(f"📊 {len(company_info['text_content']):,} characters total")
            
            # Display the COMPLETE text content
            st.markdown("#### Full Content:")
            
            # Use a text area for better readability of long content
            st.text_area(
                "Complete Text Content",
                value=company_info['text_content'],
                height=400,  # Larger height to show more content
                disabled=True,
                label_visibility="collapsed"
            )
            
            # Also show as markdown for better formatting
            st.markdown("#### Formatted Content:")
            st.markdown(company_info['text_content'])
            
        else:
            st.warning("No text content found")
        
        # Download section
        st.markdown("---")
        st.markdown("### 💾 Download Options")
        col1, col2, col3 = st.columns(3)
        with col1:
            if 'github_url' in company_info:
                st.markdown(f"[📥 Download Original DOCX]({company_info['github_url']})")
        with col2:
            # Download all links as text file
            links_text = f"# {company_info['company_name']} - Extracted Links\n\n"
            for i, link in enumerate(company_info['links'], 1):
                links_text += f"{i}. {link}\n"
            
            st.download_button(
                "📋 Download All Links",
                data=links_text,
                file_name=f"{company_info['company_name']}_links.txt",
                mime="text/plain",
                key=f"copy_links_{company_info['filename']}"
            )
        with col3:
            # Download complete text content
            full_content = f"# {company_info['company_name']} - Complete Text Content\n\n"
            full_content += f"Extraction Date: {company_info.get('extraction_date', 'Unknown')}\n"
            full_content += f"Total Links: {company_info['total_links']}\n"
            full_content += f"Text Length: {company_info['text_length']:,} characters\n\n"
            full_content += "## Complete Text:\n\n"
            full_content += company_info['text_content']
            
            st.download_button(
                "📄 Download Complete Text",
                data=full_content,
                file_name=f"{company_info['company_name']}_complete_content.txt",
                mime="text/plain",
                key=f"copy_text_{company_info['filename']}"
            )

def main():
    # Page configuration
    st.set_page_config(
        page_title="Company Research Dashboard",
        page_icon="🏢",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    # Custom CSS for better styling
    st.markdown("""
    <style>
        .main-header {
            text-align: center;
            padding: 2rem 0;
            background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
            color: white;
            border-radius: 10px;
            margin-bottom: 2rem;
        }
        .metric-card {
            background: #f8f9fa;
            padding: 1rem;
            border-radius: 8px;
            border-left: 4px solid #667eea;
        }
        .stTabs [data-baseweb="tab-list"] button [data-testid="stMarkdownContainer"] p {
            font-size: 16px;
        }
    </style>
    """, unsafe_allow_html=True)
    
    # Header
    st.markdown("""
    <div class="main-header">
        <h1>🏢 Company Research Dashboard</h1>
        <p>View extracted company data, links, and insights from DOCX files</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Sidebar
    st.sidebar.header("⚙️ Configuration")
    
    # Data source selection
    data_source = st.sidebar.selectbox(
        "Choose data source:",
        ["GitHub Repository", "Sample Data (Demo)", "Local Directory"]
    )
    
    company_data = []
    
    if data_source == "GitHub Repository":
        st.sidebar.markdown("### 🔧 GitHub Configuration")
        
        # GitHub repository URL input
        repo_url = st.sidebar.text_input(
            "GitHub Repository URL:",
            value="https://github.com/Abdulmasood14/news_ai_mode",
            help="Enter the GitHub repository URL containing DOCX files"
        )
        
        # Data folder path
        data_folder = st.sidebar.text_input(
            "Data Folder Path:",
            value="data/",
            help="Folder path within the repo containing DOCX files (include trailing slash)"
        )
        
        if repo_url:
            # Convert GitHub URL to raw URLs
            if "github.com" in repo_url:
                # Convert to raw.githubusercontent.com format
                if "/tree/" in repo_url:
                    raw_base = repo_url.replace("github.com", "raw.githubusercontent.com").replace("/tree/", "/")
                else:
                    # Handle direct repo URL
                    raw_base = repo_url.replace("github.com", "raw.githubusercontent.com")
                    if not raw_base.endswith("/main") and not raw_base.endswith("/master"):
                        if not raw_base.endswith("/"):
                            raw_base += "/"
                        raw_base += "main/"
                
                if not raw_base.endswith("/"):
                    raw_base += "/"
                
                st.sidebar.info("📝 List your DOCX filenames below (auto-filled)")
                
                # Auto-fill with all your filenames
                default_filenames = """bajaj-auto_complete_20250819_172757.docx
bajfinspv_complete_20250819_172425.docx
bajfinance_complete_20250819_173439.docx
coalindia_complete_20250819_173259.docx
drreddy_complete_20250819_173118.docx
hdfclife_complete_20250819_172108.docx
itc_complete_20250819_172939.docx
lt_complete_20250819_173935.docx
maruti_complete_20250819_173801.docx
ntpc_complete_20250819_171629.docx
ongc_complete_20250819_173618.docx
tataconsum_complete_20250819_172249.docx
tcs_complete_20250819_172615.docx
techm_complete_20250819_171759.docx
titan_complete_20250819_171455.docx
trent_complete_20250819_171936.docx"""
                
                docx_filenames = st.sidebar.text_area(
                    "DOCX Filenames (one per line):",
                    value=default_filenames,
                    height=200,
                    help="List your DOCX filenames, one per line"
                )
                
                if docx_filenames.strip():
                    filenames = [f.strip() for f in docx_filenames.split('\n') if f.strip()]
                    github_urls = [f"{raw_base}{data_folder}{filename}" for filename in filenames]
                    
                    # Show the URLs that will be used
                    st.sidebar.write("🔗 URLs to be processed:")
                    for i, url in enumerate(github_urls[:3], 1):
                        st.sidebar.caption(f"{i}. {url}")
                    if len(github_urls) > 3:
                        st.sidebar.caption(f"... and {len(github_urls)-3} more")
                    
                    if st.sidebar.button("🚀 Load All Files"):
                        with st.spinner("🔄 Loading data from GitHub..."):
                            company_data = load_docx_from_github(github_urls)
            else:
                st.sidebar.error("❌ Invalid GitHub URL format")
    
    elif data_source == "Local Directory":
        directory_path = st.sidebar.text_input(
            "Local Directory Path:",
            placeholder="/path/to/your/docx/files",
            help="Enter the full path to directory containing DOCX files"
        )
        
        if directory_path and os.path.exists(directory_path):
            with st.spinner("🔄 Loading local files..."):
                company_data = load_docx_files_from_directory(directory_path)
        elif directory_path:
            st.sidebar.error("❌ Directory not found")
    
    else:  # Sample Data
        st.info("🚀 Showing sample data with all your companies from GitHub repository")
        
        # All your companies with sample data
        companies_sample = [
            {
                'filename': 'bajaj-auto_complete_20250819_172757.docx',
                'company_name': 'Bajaj Auto',
                'links': 19,
                'text_length': 2306,
                'content': 'Bajaj Auto reported an 8% increase in total sales (including exports) in May 2025 compared to May 2024, reaching 3,84,621 units. Exports were a key driver of this growth, surging by 22% year-on-year to 1,58,888 vehicles in May 2025.'
            },
            {
                'filename': 'bajfinspv_complete_20250819_172425.docx',
                'company_name': 'Bajfinspv',
                'links': 15,
                'text_length': 1850,
                'content': 'Bajaj Finserv reported strong quarterly results with significant growth in assets under management and digital lending portfolio expansion across multiple financial services segments.'
            },
            {
                'filename': 'bajfinance_complete_20250819_173439.docx',
                'company_name': 'Bajfinance',
                'links': 22,
                'text_length': 2890,
                'content': 'Bajaj Finance continues to dominate the NBFC sector with robust loan growth, improved asset quality metrics, and strategic expansion in rural and semi-urban markets.'
            },
            {
                'filename': 'coalindia_complete_20250819_173259.docx',
                'company_name': 'Coalindia',
                'links': 18,
                'text_length': 2156,
                'content': 'Coal India Limited maintains its position as the world largest coal producer, focusing on sustainable mining practices and renewable energy initiatives.'
            },
            {
                'filename': 'drreddy_complete_20250819_173118.docx',
                'company_name': 'Drreddy',
                'links': 16,
                'text_length': 1967,
                'content': 'Dr. Reddy Laboratories announced strong performance in both domestic and international markets, with several new drug approvals and pipeline expansions.'
            },
            {
                'filename': 'hdfclife_complete_20250819_172108.docx',
                'company_name': 'Hdfclife',
                'links': 14,
                'text_length': 1745,
                'content': 'HDFC Life Insurance continues to grow its market share through innovative product offerings and digital transformation initiatives in the life insurance sector.'
            },
            {
                'filename': 'itc_complete_20250819_172939.docx',
                'company_name': 'Itc',
                'links': 20,
                'text_length': 2534,
                'content': 'ITC Limited diversifies its portfolio beyond tobacco, showing strong growth in FMCG, hotels, and agri-business segments with sustainability focus.'
            },
            {
                'filename': 'lt_complete_20250819_173935.docx',
                'company_name': 'Lt',
                'links': 17,
                'text_length': 2098,
                'content': 'Larsen & Toubro demonstrates strong execution capabilities across infrastructure, defense, and technology sectors with significant order book growth.'
            },
            {
                'filename': 'maruti_complete_20250819_173801.docx',
                'company_name': 'Maruti',
                'links': 21,
                'text_length': 2678,
                'content': 'Maruti Suzuki maintains its leadership in the Indian automobile market with strong sales performance and expansion into electric vehicle segment.'
            },
            {
                'filename': 'ntpc_complete_20250819_171629.docx',
                'company_name': 'Ntpc',
                'links': 13,
                'text_length': 1623,
                'content': 'NTPC Limited continues its transition towards renewable energy while maintaining its position as India largest power generation company.'
            },
            {
                'filename': 'ongc_complete_20250819_173618.docx',
                'company_name': 'Ongc',
                'links': 19,
                'text_length': 2289,
                'content': 'Oil and Natural Gas Corporation explores new opportunities in clean energy while optimizing production from existing oil and gas fields.'
            },
            {
                'filename': 'tataconsum_complete_20250819_172249.docx',
                'company_name': 'Tataconsum',
                'links': 12,
                'text_length': 1456,
                'content': 'Tata Consumer Products expands its portfolio through strategic acquisitions and innovation in health and wellness beverage categories.'
            },
            {
                'filename': 'tcs_complete_20250819_172615.docx',
                'company_name': 'Tcs',
                'links': 25,
                'text_length': 3124,
                'content': 'Tata Consultancy Services maintains its leadership in IT services with strong growth in digital transformation and cloud migration services.'
            },
            {
                'filename': 'techm_complete_20250819_171759.docx',
                'company_name': 'Techm',
                'links': 16,
                'text_length': 1987,
                'content': 'Tech Mahindra focuses on 5G, IoT, and digital transformation solutions while expanding its presence in emerging technology markets.'
            },
            {
                'filename': 'titan_complete_20250819_171455.docx',
                'company_name': 'Titan',
                'links': 18,
                'text_length': 2234,
                'content': 'Titan Company continues to dominate the jewelry and watches market with strong brand presence and expansion in wedding and fashion jewelry.'
            },
            {
                'filename': 'trent_complete_20250819_171936.docx',
                'company_name': 'Trent',
                'links': 14,
                'text_length': 1678,
                'content': 'Trent Limited shows robust growth in retail fashion with successful expansion of Westside and Zudio store formats across India.'
            }
        ]
        
        # Convert to proper format
        company_data = []
        for comp in companies_sample:
            # Create sample links
            sample_links = [
                f"https://economictimes.com/{comp['company_name'].lower()}-news-analysis",
                f"https://moneycontrol.com/stocks/{comp['company_name'].lower()}-share-price",
                f"https://nseindia.com/get-quotes/equity?symbol={comp['company_name'].upper()}",
                f"https://business-standard.com/topic/{comp['company_name'].lower()}",
                f"https://financialexpress.com/market/stock/{comp['company_name'].lower()}"
            ]
            
            # Add more links to match the count
            while len(sample_links) < comp['links']:
                sample_links.append(f"https://example-news-{len(sample_links)+1}.com/{comp['company_name'].lower()}")
            
            company_data.append({
                'company_name': comp['company_name'],
                'filename': comp['filename'],
                'summary': f"Total links: {comp['links']}, Text length: {comp['text_length']:,} characters, Date: 2025-08-19 17:27:57",
                'extraction_date': "2025-08-19 17:27:57",
                'links': sample_links[:comp['links']],
                'text_content': comp['content'],
                'total_links': comp['links'],
                'total_links_from_summary': comp['links'],
                'text_length': len(comp['content']),
                'text_length_from_summary': comp['text_length']
            })
    
    # Main content
    if company_data:
        # Statistics
        st.markdown("## 📊 Dashboard Overview")
        
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("🏢 Companies", len(company_data))
        with col2:
            total_links = sum(comp['total_links'] for comp in company_data)
            st.metric("🔗 Total Links", total_links)
        with col3:
            avg_links = total_links / len(company_data) if company_data else 0
            st.metric("📈 Avg Links/Company", f"{avg_links:.1f}")
        with col4:
            total_text = sum(comp['text_length'] for comp in company_data)
            st.metric("📝 Total Characters", f"{total_text:,}")
        
        st.markdown("---")
        
        # Search functionality
        search_term = st.text_input(
            "🔍 Search companies:",
            placeholder="Enter company name to filter...",
            help="Search by company name"
        )
        
        # Filter companies
        filtered_data = company_data
        if search_term:
            filtered_data = [
                comp for comp in company_data 
                if search_term.lower() in comp['company_name'].lower()
            ]
        
        # Sort options
        sort_option = st.selectbox(
            "📊 Sort by:",
            ["Company Name", "Number of Links (High to Low)", "Number of Links (Low to High)", "Text Length (High to Low)"]
        )
        
        if sort_option == "Company Name":
            filtered_data = sorted(filtered_data, key=lambda x: x['company_name'])
        elif sort_option == "Number of Links (High to Low)":
            filtered_data = sorted(filtered_data, key=lambda x: x['total_links'], reverse=True)
        elif sort_option == "Number of Links (Low to High)":
            filtered_data = sorted(filtered_data, key=lambda x: x['total_links'])
        elif sort_option == "Text Length (High to Low)":
            filtered_data = sorted(filtered_data, key=lambda x: x['text_length'], reverse=True)
        
        # Display results
        st.markdown(f"## 📋 Company Data ({len(filtered_data)} companies)")
        
        if filtered_data:
            for company_info in filtered_data:
                display_company_card(company_info)
        else:
            st.warning("🔍 No companies found matching your search criteria")
            
    else:
        st.markdown("## 🚀 Get Started")
        st.markdown("""
        ### How to use this dashboard:
        
        1. **📊 Sample Data (Demo):**
           - See how the dashboard works with your exact DOCX format
           - Perfect for testing and demonstrations
        
        2. **🌐 GitHub Repository (Recommended for sharing):**
           - Upload your DOCX files to a GitHub repository
           - Enter the repository URL in the sidebar
           - List your DOCX filenames
           - Deploy to Streamlit Cloud and share the URL!
        
        3. **💻 Local Directory:**
           - For testing with local files
           - Enter the path to your DOCX files folder
        
        ### ✅ Expected DOCX format:
        Your DOCX files should contain exactly what you showed:
        - **Summary** section with total links, text length, and date
        - **Extracted Links** section with numbered URLs
        - **Extracted Text Content** section with the main content
        
        ### 🚀 Ready to deploy?
        1. Create a GitHub repository
        2. Add your DOCX files to a `data/` folder
        3. Add this code as `app.py`
        4. Deploy on [Streamlit Cloud](https://share.streamlit.io)
        5. Share your public URL with others!
        """)

if __name__ == "__main__":
    main()
